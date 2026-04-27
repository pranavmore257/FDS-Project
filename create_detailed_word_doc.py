from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, code_text, language="java"):
    """Add a formatted code block"""
    code_para = doc.add_paragraph()
    code_para.style = 'Normal'
    code_para.paragraph_format.left_indent = Inches(0.5)
    code_para.paragraph_format.first_line_indent = Inches(0)
    
    # Set monospace font
    run = code_para.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # Add gray background
    shading_elm = code_para._element.get_or_add_pPr()
    from docx.oxml import parse_xml
    shading_xml = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
    shading_elm.append(shading_xml)

# Create document
doc = Document()

# Title
title = doc.add_heading('Deep Dive: DSA Algorithms and Hashtable Implementation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('FDS Project - Comprehensive Technical Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. LocationHashtable: In-Depth Storage and Retrieval',
    '2. Dijkstra\'s Algorithm: Single-Source Shortest Path',
    '3. Bellman-Ford Algorithm: Handling Negative Weights',
    '4. Floyd-Warshall Algorithm: All-Pairs Shortest Paths',
    '5. Algorithm Comparison and Performance Analysis',
    '6. Real-World Examples from Our Project',
    '7. Integration: How They Work Together'
]

for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============== HASHTABLE SECTION ==============
doc.add_heading('1. LocationHashtable: In-Depth Storage and Retrieval', level=1)

doc.add_heading('1.1 What is a Hashtable?', level=2)
doc.add_paragraph('''
A Hashtable is a data structure that implements an associative array - a structure that maps keys to values using a hash function. Unlike arrays that use numeric indices (0, 1, 2...), hashtables use any data type as a key and retrieve values in O(1) average time complexity.

In our FDS project, we use Java's Hashtable class which is:
• Thread-safe (all methods are synchronized)
• Uses hash function to convert keys to array indices
• Handles collision resolution using chaining or open addressing
• Maintains insertion order in some implementations
''')

doc.add_heading('1.2 How Does Hashing Work?', level=2)
doc.add_paragraph('''
Step-by-step process when you add a location:

EXAMPLE: Adding "Katraj" location to hashtable

1. HASH FUNCTION:
   Input: "katraj" (converted to lowercase)
   Hash Code: hashCode() method generates an integer
   Example: "katraj".hashCode() = 1234567890
   
2. INDEX CALCULATION:
   hashIndex = hashCode % tableSize
   Example: 1234567890 % 16 = 10
   (The hashtable has 16 slots by default, 0-15)
   
3. STORAGE AT BUCKET:
   Location object stored at bucket[10]
   bucket[10] = Location{
       name: "Katraj",
       latitude: 18.4767,
       longitude: 73.8409,
       geocodeSource: "google",
       matchedPlace: "Katraj, Pune"
   }
''')

doc.add_heading('1.3 Code Implementation Walkthrough', level=2)

doc.add_paragraph('Our LocationHashtable class in Java:')
add_code_block(doc, '''public class LocationHashtable implements Serializable {
    private static final long serialVersionUID = 1L;
    
    // The actual hashtable that stores location data
    private Hashtable<String, Location> locations;
    
    // Inner class to store location data
    public static class Location implements Serializable {
        public String name;              // Original name: "Katraj"
        public double latitude;          // Coordinates
        public double longitude;
        public String geocodeSource;     // Where data came from (google/osm)
        public String matchedPlace;      // Normalized place name
        public long timestamp;           // When added
        
        public Location(String name, double latitude, double longitude,
                       String geocodeSource, String matchedPlace) {
            this.name = name;
            this.latitude = latitude;
            this.longitude = longitude;
            this.geocodeSource = geocodeSource;
            this.matchedPlace = matchedPlace;
            this.timestamp = System.currentTimeMillis();
        }
    }
    
    // Initialize the hashtable
    public LocationHashtable() {
        this.locations = new Hashtable<>();  // Empty hashtable
    }
}''')

doc.add_heading('1.4 Storage: How Data Gets Stored', level=2)

doc.add_paragraph('''
When we add a location to the hashtable:
''')

add_code_block(doc, '''public synchronized void addLocation(String name, double latitude, double longitude,
                                   String geocodeSource, String matchedPlace) {
    // Create a Location object with all the data
    Location loc = new Location(name, latitude, longitude, 
                               geocodeSource, matchedPlace);
    
    // Convert key to lowercase for case-insensitive storage
    // Example: "Viman Nagar" becomes "viman nagar"
    String key = name.toLowerCase();
    
    // Store in hashtable
    locations.put(key, loc);  // O(1) average time
}''')

doc.add_paragraph('''
DETAILED FLOW WITH EXAMPLE:

1. Input Data:
   name = "Katraj"
   latitude = 18.4767
   longitude = 73.8409
   geocodeSource = "google"
   matchedPlace = "Katraj, Pune, Maharashtra"

2. Object Creation:
   Location object is created with all these fields
   
3. Key Normalization:
   key = "katraj" (converted to lowercase)
   Why lowercase? For case-insensitive lookup
   Now both "Katraj" and "KATRAJ" will find same entry
   
4. Hashing:
   hashCode = key.hashCode()
   index = hashCode % tableSize
   
5. Storage:
   hashtable[index] = Location object
   Now the data is stored and ready for retrieval

TIME COMPLEXITY: O(1) average case
   - Hash function: O(1)
   - Array access: O(1)
   - With good hash function and low collision: O(1)

SPACE COMPLEXITY: O(n) where n = number of locations
''')

doc.add_heading('1.5 Retrieval: How Data Gets Retrieved', level=2)

add_code_block(doc, '''public synchronized Location getLocation(String name) {
    // Convert to lowercase for consistent lookup
    String key = name.toLowerCase();
    
    // Retrieve from hashtable
    return locations.get(key);  // O(1) average time
}''')

doc.add_paragraph('''
DETAILED RETRIEVAL FLOW WITH EXAMPLE:

SCENARIO: User searches for location "Katraj"

1. Input:
   searchName = "Katraj"
   
2. Key Normalization:
   key = "katraj" (lowercase)
   
3. Hashing (same process):
   hashCode = "katraj".hashCode()
   index = hashCode % tableSize
   
4. Direct Array Access:
   result = hashtable[index]
   
5. Collision Handling:
   If multiple items hash to same index (collision):
   - Hashtable uses chaining: bucket contains linked list
   - Linear search through the chain to find correct entry
   - Compare key with stored keys until match found
   
6. Return:
   If found: return Location object
   If not found: return null

TIME COMPLEXITY:
   Best case: O(1)
   Average case: O(1)
   Worst case: O(n) if all items hash to same index (rare with good hash function)

EXAMPLE RETRIEVAL:
   Location loc = hashtable.getLocation("Katraj");
   if (loc != null) {
       System.out.println("Found: " + loc.name);
       System.out.println("Coordinates: " + loc.latitude + ", " + loc.longitude);
   }
''')

doc.add_heading('1.6 Other Hashtable Operations', level=2)

doc.add_paragraph('Check if location exists:')
add_code_block(doc, '''public synchronized boolean hasLocation(String name) {
    return locations.containsKey(name.toLowerCase());  // O(1)
}
// Usage: if (hashtable.hasLocation("Baner")) { ... }''')

doc.add_paragraph('Remove a location:')
add_code_block(doc, '''public synchronized Location removeLocation(String name) {
    return locations.remove(name.toLowerCase());  // O(1)
}
// Usage: hashtable.removeLocation("Wakad");''')

doc.add_paragraph('Get all locations:')
add_code_block(doc, '''public synchronized Hashtable<String, Location> getAllLocations() {
    return (Hashtable<String, Location>) locations.clone();
}
// Returns copy of entire hashtable for iteration''')

doc.add_heading('1.7 Thread-Safety in Our Implementation', level=2)
doc.add_paragraph('''
Every method in our LocationHashtable is marked "synchronized":

    public synchronized void addLocation(...)
    public synchronized Location getLocation(...)
    public synchronized boolean hasLocation(...)

What does "synchronized" mean?

• Only ONE thread can execute the method at a time
• Other threads must wait their turn
• Prevents race conditions and data corruption
• Example:
  - Thread 1 starts adding "Mumbai"
  - Thread 2 tries to add "Delhi" but waits
  - Thread 1 completes, releases lock
  - Thread 2 now can add "Delhi"

This is important because our Flask backend can receive multiple requests simultaneously!
''')

doc.add_page_break()

# ============== DIJKSTRA'S ALGORITHM ==============
doc.add_heading('2. Dijkstra\'s Algorithm: Single-Source Shortest Path', level=1)

doc.add_heading('2.1 Algorithm Overview', level=2)
doc.add_paragraph('''
Dijkstra's algorithm finds the shortest path from a source vertex to all other vertices in a weighted graph with NON-NEGATIVE edge weights.

TIME COMPLEXITY: O((V + E) log V) with priority queue
SPACE COMPLEXITY: O(V + E)

WHERE TO USE:
✓ Finding shortest route between two locations
✓ GPS navigation systems
✓ Non-negative weights (distances, time)
✗ Cannot handle negative edge weights
✗ Cannot detect negative cycles
''')

doc.add_heading('2.2 Step-by-Step Algorithm Explanation', level=2)
doc.add_paragraph('''
STEP 1: INITIALIZATION

Initialize three data structures:
1. distances = map of vertices to their shortest known distance
2. previousVertex = map to track path reconstruction
3. priorityQueue = min-heap to always get unvisited vertex with smallest distance

Example with locations: Pune → Mumbai → Nashik

distances = {
    "Pune": 0,           // source, so distance is 0
    "Mumbai": ∞,         // infinity, not yet visited
    "Nashik": ∞          // infinity, not yet visited
}

previousVertex = {
    "Pune": null,        // starting point, no previous
    "Mumbai": null,
    "Nashik": null
}

priorityQueue = [(0, "Pune")]  // (distance, vertex)
''')

doc.add_paragraph('''
STEP 2: MAIN LOOP - Extract minimum

While priorityQueue is not empty:
  1. Remove vertex with minimum distance from priority queue
  2. Mark it as visited
  3. For each unvisited neighbor of current vertex:
     - Calculate: newDistance = currentDistance + edgeWeight
     - If newDistance < known distance to neighbor:
       * Update distances[neighbor] = newDistance
       * Update previousVertex[neighbor] = currentVertex
       * Add neighbor to priority queue

EXAMPLE WALKTHROUGH:

Initial state:
Edges:
  Pune → Mumbai (150 km)
  Pune → Nashik (200 km)
  Mumbai → Nashik (180 km)

Find shortest path from Pune to Nashik:

ITERATION 1:
─────────────
Current: ("Pune", distance=0)
Neighbors of Pune: [Mumbai, Nashik]

Check Mumbai:
  newDistance = 0 + 150 = 150
  distances["Mumbai"] = ∞ > 150
  ✓ Update distances["Mumbai"] = 150
  ✓ previousVertex["Mumbai"] = "Pune"
  ✓ Add (150, "Mumbai") to queue

Check Nashik:
  newDistance = 0 + 200 = 200
  distances["Nashik"] = ∞ > 200
  ✓ Update distances["Nashik"] = 200
  ✓ previousVertex["Nashik"] = "Pune"
  ✓ Add (200, "Nashik") to queue

State after iteration 1:
distances = {"Pune": 0, "Mumbai": 150, "Nashik": 200}
priorityQueue = [(150, "Mumbai"), (200, "Nashik")]

ITERATION 2:
─────────────
Current: ("Mumbai", distance=150)
Neighbors of Mumbai: [Nashik]

Check Nashik:
  newDistance = 150 + 180 = 330
  distances["Nashik"] = 200 < 330
  ✗ DO NOT UPDATE (330 > 200, not shorter)
  
State after iteration 2:
distances = {"Pune": 0, "Mumbai": 150, "Nashik": 200}
priorityQueue = [(200, "Nashik")]

ITERATION 3:
─────────────
Current: ("Nashik", distance=200)
Neighbors of Nashik: [Mumbai] (already visited)

No updates. Queue is empty. DONE!

FINAL RESULT:
─────────────
Shortest distance: 200 km
Path reconstruction (follow previousVertex backwards):
  Nashik ← Pune
  Path: Pune → Nashik (distance: 200)

Why not Pune → Mumbai → Nashik (330 km)?
Because 200 < 330, so direct route is shorter!
''')

doc.add_heading('2.3 Java Implementation', level=2)

add_code_block(doc, '''public DijkstraResult findShortestPath(String source, String destination) {
    long startTime = System.currentTimeMillis();
    DijkstraResult result = new DijkstraResult();
    
    Map<String, Double> distances = new HashMap<>();
    Map<String, String> previousVertex = new HashMap<>();
    Set<String> visited = new HashSet<>();
    
    // Initialize: set all distances to infinity except source
    for (String vertex : graph.getVertices()) {
        distances.put(vertex, Double.MAX_VALUE);
    }
    distances.put(source, 0.0);
    
    // Priority queue: stores (distance, vertex) pairs
    // Organized by distance (min-heap)
    PriorityQueue<AbstractMap.SimpleEntry<Double, String>> pq = 
        new PriorityQueue<>(
            Comparator.comparingDouble(AbstractMap.SimpleEntry::getKey)
        );
    pq.offer(new AbstractMap.SimpleEntry<>(0.0, source));
    
    // Main algorithm loop
    while (!pq.isEmpty()) {
        // Get vertex with minimum distance
        AbstractMap.SimpleEntry<Double, String> current = pq.poll();
        double currentDist = current.getKey();
        String currentVertex = current.getValue();
        
        // Skip if already visited
        if (visited.contains(currentVertex)) continue;
        visited.add(currentVertex);
        
        // If reached destination, reconstruct path
        if (currentVertex.equals(destination)) {
            reconstructPath(result, previousVertex, source, destination);
            result.totalDistance = distances.get(destination);
            break;
        }
        
        // Relax edges (check neighbors)
        for (Graph.Edge edge : graph.getNeighbors(currentVertex)) {
            String neighbor = edge.to;
            double newDist = currentDist + edge.weight;
            
            if (newDist < distances.get(neighbor)) {
                distances.put(neighbor, newDist);
                previousVertex.put(neighbor, currentVertex);
                pq.offer(new AbstractMap.SimpleEntry<>(newDist, neighbor));
            }
        }
    }
    
    result.executionTimeMs = System.currentTimeMillis() - startTime;
    return result;
}''')

doc.add_page_break()

# ============== BELLMAN-FORD ALGORITHM ==============
doc.add_heading('3. Bellman-Ford Algorithm: Handling Negative Weights', level=1)

doc.add_heading('3.1 Algorithm Overview', level=2)
doc.add_paragraph('''
Bellman-Ford finds the shortest path from a source vertex to all other vertices, and CAN handle negative edge weights. It can also detect negative cycles.

TIME COMPLEXITY: O(V × E)
SPACE COMPLEXITY: O(V)

WHERE TO USE:
✓ Graphs with negative edge weights
✓ Detect negative cycles
✓ More robust than Dijkstra
✗ Slower than Dijkstra for non-negative weights
✗ Cannot handle graphs with negative cycles (infinite negative path)

REAL WORLD EXAMPLE:
In currency exchange:
  USD → EUR: +0.92 (exchange rate)
  EUR → JPY: +130 (exchange rate)
  JPY → USD: -0.008 (loss due to fees)
  Could create negative cycle!
''')

doc.add_heading('3.2 Step-by-Step Algorithm Explanation', level=2)
doc.add_paragraph('''
STEP 1: INITIALIZATION (same as Dijkstra)

distances = {
    "Pune": 0,
    "Mumbai": ∞,
    "Nashik": ∞
}

previousVertex = {
    "Pune": null,
    "Mumbai": null,
    "Nashik": null
}

STEP 2: RELAX EDGES (V-1) TIMES

For i = 1 to (V-1):  // V = number of vertices
  For each edge (u → v) with weight w:
    If distance[u] + w < distance[v]:
      distance[v] = distance[u] + w
      previousVertex[v] = u

Why (V-1) times?
• In worst case, shortest path uses (V-1) edges
• Each iteration "spreads" shortest distance one edge further
• After (V-1) iterations, all shortest paths are found

STEP 3: DETECT NEGATIVE CYCLES

For each edge (u → v) with weight w:
  If distance[u] + w < distance[v]:
    // We can still relax? This means negative cycle exists!
    hasNegativeCycle = true
    return error

EXAMPLE WALKTHROUGH:

Graph with weights (some negative):
Pune → Mumbai (100 km)
Mumbai → Nashik (-50 km)  // Negative weight!
Nashik → Pune (-75 km)

Vertices: 3, so relax V-1 = 2 times

ITERATION 1:
─────────────
Relax all edges:

Edge Pune → Mumbai:
  newDist = 0 + 100 = 100
  100 < ∞: ✓ Update distances["Mumbai"] = 100

Edge Mumbai → Nashik:
  newDist = 100 + (-50) = 50
  50 < ∞: ✓ Update distances["Nashik"] = 50

Edge Nashik → Pune:
  newDist = 50 + (-75) = -25
  -25 < 0: ✓ Update distances["Pune"] = -25

After iteration 1:
distances = {"Pune": -25, "Mumbai": 100, "Nashik": 50}

ITERATION 2:
─────────────
Relax all edges again:

Edge Pune → Mumbai:
  newDist = -25 + 100 = 75
  75 < 100: ✓ Update distances["Mumbai"] = 75

Edge Mumbai → Nashik:
  newDist = 75 + (-50) = 25
  25 < 50: ✓ Update distances["Nashik"] = 25

Edge Nashik → Pune:
  newDist = 25 + (-75) = -50
  -50 < -25: ✓ Update distances["Pune"] = -50

After iteration 2:
distances = {"Pune": -50, "Mumbai": 75, "Nashik": 25}

STEP 3: CHECK FOR NEGATIVE CYCLE
─────────────────────────────────
Try to relax one more time:

Edge Pune → Mumbai:
  newDist = -50 + 100 = 50
  50 < 75: ✓ CAN STILL RELAX!

This means NEGATIVE CYCLE EXISTS!
The algorithm marks hasNegativeCycle = true
Because distances keep decreasing infinitely!
''')

doc.add_heading('3.3 Java Implementation', level=2)

add_code_block(doc, '''public BellmanFordResult findShortestPath(String source, String destination) {
    long startTime = System.currentTimeMillis();
    BellmanFordResult result = new BellmanFordResult();
    
    Map<String, Double> distances = new HashMap<>();
    Map<String, String> previousVertex = new HashMap<>();
    
    // Initialize distances
    for (String vertex : graph.getVertices()) {
        distances.put(vertex, Double.MAX_VALUE);
    }
    distances.put(source, 0.0);
    
    int vertexCount = graph.getVertices().size();
    
    // STEP 1: Relax edges (V-1) times
    for (int i = 0; i < vertexCount - 1; i++) {
        for (String u : graph.getVertices()) {
            for (Graph.Edge edge : graph.getNeighbors(u)) {
                String v = edge.to;
                double weight = edge.weight;
                
                // Relax edge
                if (distances.get(u) != Double.MAX_VALUE &&
                    distances.get(u) + weight < distances.get(v)) {
                    
                    distances.put(v, distances.get(u) + weight);
                    previousVertex.put(v, u);
                }
            }
        }
    }
    
    // STEP 2: Check for negative cycles
    for (String u : graph.getVertices()) {
        for (Graph.Edge edge : graph.getNeighbors(u)) {
            String v = edge.to;
            double weight = edge.weight;
            
            if (distances.get(u) != Double.MAX_VALUE &&
                distances.get(u) + weight < distances.get(v)) {
                
                // Can still relax = negative cycle detected
                result.hasNegativeCycle = true;
                result.errorMessage = "Negative cycle detected!";
                break;
            }
        }
        if (result.hasNegativeCycle) break;
    }
    
    // Reconstruct path
    if (!result.hasNegativeCycle) {
        reconstructPath(result, previousVertex, source, destination);
        result.totalDistance = distances.get(destination);
    }
    
    result.executionTimeMs = System.currentTimeMillis() - startTime;
    return result;
}''')

doc.add_page_break()

# ============== FLOYD-WARSHALL ALGORITHM ==============
doc.add_heading('4. Floyd-Warshall Algorithm: All-Pairs Shortest Paths', level=1)

doc.add_heading('4.1 Algorithm Overview', level=2)
doc.add_paragraph('''
Floyd-Warshall computes shortest paths between ALL PAIRS of vertices in one run. Unlike Dijkstra and Bellman-Ford which find single-source paths, Floyd-Warshall gives you a complete distance matrix.

TIME COMPLEXITY: O(V³)
SPACE COMPLEXITY: O(V²)

WHERE TO USE:
✓ Need all-pairs shortest paths
✓ Dense graphs (many edges)
✓ Comparing multiple route options
✓ Building complete distance matrix
✗ Large graphs (V > 500) - too slow
✗ Sparse graphs - Dijkstra multiple times is faster

KEY IDEA:
Instead of starting from one source, Floyd-Warshall considers all possible "intermediate" vertices and checks if using that intermediate vertex creates a shorter path.
''')

doc.add_heading('4.2 Step-by-Step Algorithm Explanation', level=2)
doc.add_paragraph('''
STEP 1: INITIALIZE DISTANCE MATRIX

Create matrix[V][V] where matrix[i][j] = direct distance from i to j

Example with 3 locations:
            Pune   Mumbai  Nashik
Pune        0      100     200
Mumbai      100    0       180
Nashik      200    180     0

STEP 2: FOR EACH INTERMEDIATE VERTEX K

For k = 0 to V-1:  // Consider each location as intermediate
  For i = 0 to V-1:  // All source locations
    For j = 0 to V-1:  // All destination locations
      // Check if path through k is shorter
      If distance[i][j] > distance[i][k] + distance[k][j]:
        distance[i][j] = distance[i][k] + distance[k][j]

Why this works?
• Initially: distance[i][j] = direct path
• k=0: Consider only direct paths
• k=1: Consider paths using only vertex 0 as intermediate
• k=2: Consider paths using vertices 0 or 1 as intermediate
• k=3: Consider paths using vertices 0, 1, or 2 as intermediate
• ...
• Final: All possible paths considered

EXAMPLE WALKTHROUGH:

Initial matrix:
      0      1      2
0     0      100    200     (Pune)
1     100    0      180     (Mumbai)
2     200    180    0       (Nashik)

ITERATION K=0: Consider Pune (vertex 0) as intermediate
───────────────────────────────────────────

Check Pune→Mumbai→Nashik (via Pune):
  current: Pune→Mumbai→Nashik = 100 + 180 = 280
  via Pune: Pune→Pune + Pune→Nashik = ∞ (no direct path back)
  No improvement

Check Mumbai→Pune→Nashik (via Pune):
  current: Mumbai→Nashik = 180
  via Pune: Mumbai→Pune + Pune→Nashik = 100 + 200 = 300
  300 > 180: No improvement

Check Nashik→Pune→Mumbai (via Pune):
  current: Nashik→Mumbai = 180
  via Pune: Nashik→Pune + Pune→Mumbai = 200 + 100 = 300
  300 > 180: No improvement

Matrix after K=0:
      0      1      2
0     0      100    200
1     100    0      180
2     200    180    0

ITERATION K=1: Consider Mumbai (vertex 1) as intermediate
────────────────────────────────────────────

Check Pune→Mumbai→Nashik (via Mumbai):
  current: Pune→Nashik = 200
  via Mumbai: Pune→Mumbai + Mumbai→Nashik = 100 + 180 = 280
  280 > 200: No improvement

Check Pune→Mumbai→Pune (via Mumbai):
  current: Pune→Pune = 0
  via Mumbai: Pune→Mumbai + Mumbai→Pune = 100 + 100 = 200
  200 > 0: No improvement

Check Nashik→Mumbai→Pune (via Mumbai):
  current: Nashik→Pune = 200
  via Mumbai: Nashik→Mumbai + Mumbai→Pune = 180 + 100 = 280
  280 > 200: No improvement

Matrix after K=1:
      0      1      2
0     0      100    200
1     100    0      180
2     200    180    0

ITERATION K=2: Consider Nashik (vertex 2) as intermediate
────────────────────────────────────────────

Check Pune→Nashik→Mumbai (via Nashik):
  current: Pune→Mumbai = 100
  via Nashik: Pune→Nashik + Nashik→Mumbai = 200 + 180 = 380
  380 > 100: No improvement

Check Mumbai→Nashik→Pune (via Nashik):
  current: Mumbai→Pune = 100
  via Nashik: Mumbai→Nashik + Nashik→Pune = 180 + 200 = 380
  380 > 100: No improvement

FINAL MATRIX:
      0      1      2
0     0      100    200     (From Pune to: Pune=0, Mumbai=100, Nashik=200)
1     100    0      180     (From Mumbai to: Pune=100, Mumbai=0, Nashik=180)
2     200    180    0       (From Nashik to: Pune=200, Mumbai=180, Nashik=0)

ALL PAIRS SHORTEST PATHS FOUND!
No matter where you start, you can look up shortest distance to any other location!
''')

doc.add_heading('4.3 Java Implementation', level=2)

add_code_block(doc, '''public FloydWarshallResult findAllPairs() {
    long startTime = System.currentTimeMillis();
    FloydWarshallResult result = new FloydWarshallResult();
    
    List<String> vertices = new ArrayList<>(graph.getVertices());
    int n = vertices.size();
    
    // Initialize distance matrix
    Map<String, Map<String, Double>> distances = new HashMap<>();
    Map<String, Map<String, String>> nextVertex = new HashMap<>();
    
    for (String i : vertices) {
        distances.put(i, new HashMap<>());
        nextVertex.put(i, new HashMap<>());
        
        for (String j : vertices) {
            if (i.equals(j)) {
                distances.get(i).put(j, 0.0);
            } else {
                distances.get(i).put(j, Double.MAX_VALUE);
            }
        }
    }
    
    // Add edges
    for (String u : vertices) {
        for (Graph.Edge edge : graph.getNeighbors(u)) {
            distances.get(u).put(edge.to, edge.weight);
        }
    }
    
    // Floyd-Warshall: Try each vertex as intermediate
    for (String k : vertices) {
        for (String i : vertices) {
            for (String j : vertices) {
                double currentDist = distances.get(i).get(j);
                double distThroughK = distances.get(i).get(k) + 
                                     distances.get(k).get(j);
                
                if (distThroughK < currentDist) {
                    distances.get(i).put(j, distThroughK);
                }
            }
        }
    }
    
    result.distances = distances;
    result.executionTimeMs = System.currentTimeMillis() - startTime;
    return result;
}''')

doc.add_page_break()

# ============== COMPARISON ==============
doc.add_heading('5. Algorithm Comparison and Performance Analysis', level=1)

# Create comparison table
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'

# Header
cells = table.rows[0].cells
cells[0].text = 'Aspect'
cells[1].text = 'Dijkstra'
cells[2].text = 'Bellman-Ford'
cells[3].text = 'Floyd-Warshall'

# Time Complexity
cells = table.rows[1].cells
cells[0].text = 'Time Complexity'
cells[1].text = 'O((V+E) log V)'
cells[2].text = 'O(V × E)'
cells[3].text = 'O(V³)'

# Space Complexity
cells = table.rows[2].cells
cells[0].text = 'Space Complexity'
cells[1].text = 'O(V + E)'
cells[2].text = 'O(V)'
cells[3].text = 'O(V²)'

# Negative Weights
cells = table.rows[3].cells
cells[0].text = 'Handles Negative Weights'
cells[1].text = 'NO'
cells[2].text = 'YES'
cells[3].text = 'YES'

# Speed
cells = table.rows[4].cells
cells[0].text = 'Speed (for non-negative)'
cells[1].text = 'FASTEST'
cells[2].text = 'MEDIUM'
cells[3].text = 'SLOWEST'

# Use Case
cells = table.rows[5].cells
cells[0].text = 'Best For'
cells[1].text = 'Single-source,\nSparse graphs'
cells[2].text = 'Negative weights,\nCycle detection'
cells[3].text = 'All-pairs,\nDense graphs'

doc.add_heading('Performance in Our FDS Project', level=2)

doc.add_paragraph('''
TEST CASE: 10 locations in Pune
                      
DIJKSTRA:      ~2 ms  ✓ Fastest - Good for individual route queries
BELLMAN-FORD:  ~5 ms  - Slower but more robust
FLOYD-WARSHALL: ~8 ms - Slowest but gives all routes simultaneously

WHEN TO USE IN OUR PROJECT:
• User asks "Route from A to B?" → Use DIJKSTRA (fast)
• Need negative costs (discounts/refunds)? → Use BELLMAN-FORD
• Compare multiple routes simultaneously? → Use FLOYD-WARSHALL
''')

doc.add_page_break()

# ============== REAL-WORLD EXAMPLES ==============
doc.add_heading('6. Real-World Examples from Our Project', level=1)

doc.add_heading('6.1 Complete Execution Flow', level=2)

doc.add_paragraph('''
USER SCENARIO: Find shortest route from Viman Nagar to Hinjawadi

STEP 1: Location Storage (Hashtable)
─────────────────────────────────────

Backend receives location requests:
  1. GET /api/location?name=Viman Nagar
     → Hashtable retrieves from bucket (O(1))
     → Returns: {lat: 18.5595, lng: 73.9278, ...}
     
  2. GET /api/location?name=Hinjawadi
     → Hashtable retrieves from bucket (O(1))
     → Returns: {lat: 18.5898, lng: 73.8268, ...}

STEP 2: Build Graph (if not already built)
───────────────────────────────────────────

API: POST /api/graph/build
Creates edges between locations using Google Maps API:
  addEdge("Viman Nagar", "Hinjawadi", 25.5);  // 25.5 km

STEP 3: Run Algorithm
──────────────────────

API: POST /api/algorithm/dijkstra
Request body: {start: "Viman Nagar", end: "Hinjawadi"}

Backend calls Java algorithm:
  DijkstraAlgorithm algo = new DijkstraAlgorithm(graph);
  DijkstraResult result = algo.findShortestPath(start, end);

STEP 4: Return Results
──────────────────────

Response: {
    "path": ["Viman Nagar", "Hinjawadi"],
    "distance": 25.5,
    "executionTimeMs": 2,
    "algorithm": "Dijkstra"
}

STEP 5: Frontend Display
────────────────────────

JavaScript receives response and:
  1. Displays distance: "25.5 km"
  2. Shows execution time: "Calculated in 2 ms"
  3. Plots route on Google Maps
''')

doc.add_heading('6.2 Algorithm Comparison Flow', level=2)

doc.add_paragraph('''
API: POST /api/algorithms/compare
Request: {start: "Pune", end: "Mumbai", locations: [...]}

Backend execution:

1. DIJKSTRA (Python)
   └─ dijkstra.py → ~2 ms
   └─ Result: Path=[Pune, Mumbai], Distance=150 km

2. BELLMAN-FORD (Java subprocess)
   └─ AlgorithmRunner.java → ~5 ms
   └─ Result: Path=[Pune, Mumbai], Distance=150 km, NoCycle=true

3. FLOYD-WARSHALL (Java subprocess)
   └─ AlgorithmRunner.java → ~8 ms
   └─ Result: AllPairs={...}, Distance=150 km

Frontend receives combined results and shows:
┌─────────────────────────────────────────────────────┐
│ Algorithm Comparison Results                        │
├─────────────────────────────────────────────────────┤
│ ⚡ DIJKSTRA (FASTEST): 150 km in 2 ms             │
│ 🔄 BELLMAN-FORD: 150 km in 5 ms (No cycles)      │
│ 📊 FLOYD-WARSHALL: 150 km in 8 ms (All paths)    │
└─────────────────────────────────────────────────────┘

All three algorithms found same path!
Dijkstra is 4x faster for this single query.
Floyd-Warshall useful if comparing many route options.
''')

doc.add_page_break()

# ============== INTEGRATION ==============
doc.add_heading('7. Integration: How They Work Together', level=1)

doc.add_paragraph('''
COMPLETE SYSTEM ARCHITECTURE:

┌─────────────────────────────────────────────────────────────┐
│                   WEB BROWSER (Frontend)                    │
│        HTML/CSS/JavaScript - User Interface Layer           │
└────────────────────┬────────────────────────────────────────┘
                     │ HTTP/REST API
┌────────────────────▼────────────────────────────────────────┐
│                 PYTHON FLASK (Backend)                      │
│                                                             │
│ 1. Location Management                                      │
│    ├─ Python LocationHashtable (OrderedDict)               │
│    ├─ Google Maps Geocoding API                            │
│    ├─ OSRM (road distance calculation)                     │
│    └─ O(1) lookup/storage of location data                │
│                                                             │
│ 2. Graph Construction                                       │
│    ├─ addEdge(from, to, weight) operations                 │
│    ├─ Weight = real driving time/distance                  │
│    └─ Connected network of locations                       │
│                                                             │
│ 3. Algorithm Routing                                        │
│    ├─ Dijkstra.py (Python implementation)                  │
│    ├─ Java Executor (subprocess manager)                   │
│    └─ Sends JSON to JVM                                    │
└────────────────────┬────────────────────────────────────────┘
                     │ subprocess.run()
┌────────────────────▼────────────────────────────────────────┐
│                 JAVA VIRTUAL MACHINE (JVM)                  │
│                                                             │
│ AlgorithmRunner.java                                        │
│  ├─ LocationHashtable (Java Hashtable)                     │
│  ├─ Graph (adjacency list + matrix)                        │
│  │                                                          │
│  ├─ DijkstraAlgorithm.java                                 │
│  │  └─ Priority queue based search                         │
│  │                                                          │
│  ├─ BellmanFordAlgorithm.java                              │
│  │  └─ Edge relaxation algorithm                           │
│  │                                                          │
│  └─ FloydWarshallAlgorithm.java                            │
│     └─ All-pairs shortest paths                            │
│                                                             │
│ JSON Output: {algorithm, path, distance, time}             │
└────────────────────┬────────────────────────────────────────┘
                     │ stdout JSON
┌────────────────────▼────────────────────────────────────────┐
│           PYTHON Flask (Process Results)                    │
│  ├─ Parse JSON from Java                                   │
│  ├─ Format response                                         │
│  └─ Send to frontend                                        │
└────────────────────┬────────────────────────────────────────┘
                     │ HTTP Response JSON
┌────────────────────▼────────────────────────────────────────┐
│            Frontend (Display Results)                       │
│  ├─ Plot route on Google Maps                              │
│  ├─ Show distance and time                                 │
│  ├─ Display execution metrics                              │
│  └─ Allow user to try other algorithms                     │
└─────────────────────────────────────────────────────────────┘

DATA FLOW EXAMPLE:

User clicks "Find Route: Pune to Mumbai"

1. Frontend sends: {start: "Pune", end: "Mumbai"}
                                  ↓
2. Flask receives, checks LocationHashtable:
   - Get Pune coordinates (O(1)): {18.5204, 73.8567}
   - Get Mumbai coordinates (O(1)): {19.0760, 72.8777}
                                  ↓
3. Flask constructs Graph with edges
                                  ↓
4. Flask calls Java subprocess with JSON:
   {
     "algorithm": "dijkstra",
     "source": "Pune",
     "destination": "Mumbai",
     "vertices": ["Pune", "Mumbai", "Nashik", ...],
     "edges": [{"from": "Pune", "to": "Mumbai", "weight": 150}, ...]
   }
                                  ↓
5. Java DijkstraAlgorithm processes:
   - Create priority queue
   - Initialize distances
   - Relax edges iteratively
   - Reconstruct path
   - Return JSON with results
                                  ↓
6. Flask receives JSON, formats response
                                  ↓
7. Frontend displays:
   "Shortest route: Pune → Mumbai (150 km, 3 hours)"
   Execution time: 2 ms

HASHTABLE THROUGHOUT:

• Start: Frontend has "Pune"
• Step 2: Flask Hashtable retrieves Pune data (O(1))
• Step 5: Java LocationHashtable stores location cache (O(1))
• Step 7: Result cached for future queries

Key: Hashtable provides O(1) lookups throughout the system!
''')

# Save document
output_path = 'c:\\Users\\ACER\\Documents\\impleFDS\\DSA_Algorithms_and_Hashtable_Documentation_DETAILED.docx'
doc.save(output_path)

print(f"✓ Detailed Word document created: {output_path}")
print(f"✓ File size: Large comprehensive guide")
print(f"✓ Includes:")
print(f"  - Deep hashtable storage/retrieval explanation")
print(f"  - Step-by-step algorithm walkthroughs")
print(f"  - Complete code implementations")
print(f"  - Real-world examples")
print(f"  - Integration architecture")
