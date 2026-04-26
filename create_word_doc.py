import docx
from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()

# Add title
doc.add_heading('DSA Algorithms and Hashtable Implementation in FDS Project', 0)

# Add table of contents placeholder
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Hashtable Implementation',
    '3. Graph Data Structure',
    '4. Shortest Path Algorithms',
    '5. Algorithm Integration',
    '6. Performance Analysis',
    '7. Real-World Examples'
]

for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Add Project Overview
doc.add_heading('Project Overview', level=1)
doc.add_paragraph('''
The FDS (File System and Data Structures) Project implements a comprehensive route finding system that demonstrates practical applications of fundamental Data Structures and Algorithms (DSA). The project combines Java-based algorithm implementations with a Python Flask backend to create a robust location-based service.
''')

doc.add_heading('Key Components:', level=2)
doc.add_paragraph('• LocationHashtable: Persistent storage for location data')
doc.add_paragraph('• Graph Data Structure: Adjacency list and matrix representations')
doc.add_paragraph('• Three Shortest Path Algorithms: Dijkstra, Bellman-Ford, Floyd-Warshall')
doc.add_paragraph('• Hybrid Architecture: Java algorithms + Python backend')

# Add Hashtable Implementation
doc.add_heading('Hashtable Implementation', level=1)
doc.add_heading('Java Implementation (LocationHashtable.java)', level=2)
doc.add_paragraph('The project uses Java\'s built-in Hashtable class for thread-safe location storage:')

# Add code block for Java class
code_para = doc.add_paragraph()
code_para.add_run('public class LocationHashtable implements Serializable {').bold = True
code_para.add_run('\n    private Hashtable<String, Location> locations;')
code_para.add_run('\n    \n')
code_para.add_run('    public static class Location implements Serializable {').bold = True
code_para.add_run('\n        public String name;')
code_para.add_run('\n        public double latitude;')
code_para.add_run('\n        public double longitude;')
code_para.add_run('\n        public String geocodeSource;')
code_para.add_run('\n        public String matchedPlace;')
code_para.add_run('\n        public long timestamp;')
code_para.add_run('\n    }')
code_para.add_run('\n}')

doc.add_heading('Key Hashtable Operations:', level=2)
doc.add_paragraph('1. Add Location:')
code_para = doc.add_paragraph()
code_para.add_run('public synchronized void addLocation(String name, double latitude, double longitude,').bold = True
code_para.add_run('\n                                   String geocodeSource, String matchedPlace) {')
code_para.add_run('\n    Location loc = new Location(name, latitude, longitude, geocodeSource, matchedPlace);')
code_para.add_run('\n    locations.put(name.toLowerCase(), loc);  // Case-insensitive key')
code_para.add_run('\n}')

doc.add_paragraph('2. Retrieve Location:')
code_para = doc.add_paragraph()
code_para.add_run('public synchronized Location getLocation(String name) {').bold = True
code_para.add_run('\n    return locations.get(name.toLowerCase());')
code_para.add_run('\n}')

doc.add_heading('Hashtable Benefits in This Project:', level=2)
doc.add_paragraph('• O(1) Average Time Complexity: Fast lookups for location data')
doc.add_paragraph('• Thread Safety: Synchronized methods prevent concurrent access issues')
doc.add_paragraph('• Case Insensitivity: All keys converted to lowercase for consistent access')
doc.add_paragraph('• Serialization: Supports persistent storage across application restarts')

# Add Graph Data Structure
doc.add_heading('Graph Data Structure', level=1)
doc.add_heading('Dual Representation Approach', level=2)
doc.add_paragraph('The project implements a hybrid graph representation to support different algorithms efficiently:')

code_para = doc.add_paragraph()
code_para.add_run('public class Graph {').bold = True
code_para.add_run('\n    private Map<String, List<Edge>> adjacencyList;  // For Dijkstra/Bellman-Ford')
code_para.add_run('\n    private Set<String> vertices;')
code_para.add_run('\n    private Map<String, Map<String, Double>> adjMatrix;  // For Floyd-Warshall')
code_para.add_run('\n    private boolean isDirected;')
code_para.add_run('\n}')

doc.add_heading('Graph Construction Example:', level=2)
code_para = doc.add_paragraph()
code_para.add_run('// Adding edges to the graph').bold = True
code_para.add_run('\ngraph.addEdge("Pune", "Mumbai", 150.5);  // 150.5 km')
code_para.add_run('\ngraph.addEdge("Mumbai", "Nashik", 165.2);')
code_para.add_run('\ngraph.addEdge("Pune", "Nashik", 210.8);')

doc.add_heading('Why Dual Representation?', level=2)
doc.add_paragraph('• Adjacency List: Efficient for sparse graphs (O(V + E) space)')
doc.add_paragraph('• Adjacency Matrix: Required for Floyd-Warshall algorithm (O(V²) space)')
doc.add_paragraph('• Flexibility: Supports both directed and undirected graphs')

# Add Shortest Path Algorithms
doc.add_heading('Shortest Path Algorithms', level=1)
doc.add_heading('1. Dijkstra\'s Algorithm', level=2)
doc.add_paragraph('Purpose: Find shortest path from source to all vertices with non-negative weights.')
doc.add_paragraph('Time Complexity: O((V + E) log V) using priority queue')
doc.add_paragraph('When to Use: Graphs with non-negative edge weights, Single-source shortest path, Need fastest performance for sparse graphs')

doc.add_heading('2. Bellman-Ford Algorithm', level=2)
doc.add_paragraph('Purpose: Find shortest path from source to all vertices, handles negative weights.')
doc.add_paragraph('Time Complexity: O(V × E)')
doc.add_paragraph('When to Use: Graphs with negative edge weights, Need to detect negative cycles, Single-source shortest path with negative weights')

doc.add_heading('3. Floyd-Warshall Algorithm', level=2)
doc.add_paragraph('Purpose: Find shortest paths between all pairs of vertices.')
doc.add_paragraph('Time Complexity: O(V³)')
doc.add_paragraph('When to Use: Need all-pairs shortest paths, Dense graphs, Graphs with negative weights (but no negative cycles)')

# Add Performance Analysis
doc.add_heading('Performance Analysis', level=1)
doc.add_heading('Algorithm Comparison', level=2)

# Add table
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Algorithm'
hdr_cells[1].text = 'Time Complexity'
hdr_cells[2].text = 'Space Complexity'
hdr_cells[3].text = 'Negative Weights'
hdr_cells[4].text = 'Use Case'

row_cells = table.rows[1].cells
row_cells[0].text = 'Dijkstra'
row_cells[1].text = 'O((V + E) log V)'
row_cells[2].text = 'O(V)'
row_cells[3].text = 'No'
row_cells[4].text = 'Single-source, fastest'

row_cells = table.rows[2].cells
row_cells[0].text = 'Bellman-Ford'
row_cells[1].text = 'O(V × E)'
row_cells[2].text = 'O(V)'
row_cells[3].text = 'Yes'
row_cells[4].text = 'Single-source with negatives'

row_cells = table.rows[3].cells
row_cells[0].text = 'Floyd-Warshall'
row_cells[1].text = 'O(V³)'
row_cells[2].text = 'O(V²)'
row_cells[3].text = 'Yes'
row_cells[4].text = 'All-pairs shortest path'

# Add Real-World Examples
doc.add_heading('Real-World Examples', level=1)
doc.add_heading('Example 1: Pune-Mumbai Route Finding', level=2)
doc.add_paragraph('Input Data:')
code_para = doc.add_paragraph()
code_para.add_run('{').bold = True
code_para.add_run('\n  "vertices": ["Pune", "Mumbai", "Nashik", "Lonavala"],')
code_para.add_run('\n  "edges": [')
code_para.add_run('\n    {"from": "Pune", "to": "Mumbai", "weight": 150.5},')
code_para.add_run('\n    {"from": "Pune", "to": "Lonavala", "weight": 64.2},')
code_para.add_run('\n    {"from": "Lonavala", "to": "Mumbai", "weight": 83.1}')
code_para.add_run('\n  ],')
code_para.add_run('\n  "source": "Pune",')
code_para.add_run('\n  "destination": "Mumbai"')
code_para.add_run('\n}')

doc.add_paragraph('Dijkstra Result: Path = ["Pune", "Lonavala", "Mumbai"], Distance = 147.3 km')

doc.add_heading('Example 2: Location Hashtable Usage', level=2)
doc.add_paragraph('Adding Locations:')
code_para = doc.add_paragraph()
code_para.add_run('LocationHashtable hashtable = new LocationHashtable();').bold = True
code_para.add_run('\nhashtable.addLocation("Pune", 18.5204, 73.8567, "Google Maps", "Pune, Maharashtra");')
code_para.add_run('\nhashtable.addLocation("Mumbai", 19.0760, 72.8777, "Google Maps", "Mumbai, Maharashtra");')

# Add Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph('''
This project demonstrates a sophisticated implementation of fundamental DSA concepts in a real-world application. The combination of hashtable-based location storage, multiple graph representations, and three different shortest path algorithms provides a comprehensive learning platform for understanding:
''')

doc.add_paragraph('• Data Structure Selection: Choosing the right structure for specific use cases')
doc.add_paragraph('• Algorithm Trade-offs: Understanding when to use each algorithm')
doc.add_paragraph('• System Integration: Combining different programming languages effectively')
doc.add_paragraph('• Performance Optimization: Implementing efficient solutions with proper complexity analysis')

# Add file locations
doc.add_heading('File Locations in Project:', level=2)
doc.add_paragraph('• FDSproject/backend/algorithms/LocationHashtable.java - Main hashtable implementation')
doc.add_paragraph('• FDSproject/backend/algorithms/Graph.java - Graph data structure')
doc.add_paragraph('• FDSproject/backend/algorithms/DijkstraAlgorithm.java - Dijkstra implementation')
doc.add_paragraph('• FDSproject/backend/algorithms/BellmanFordAlgorithm.java - Bellman-Ford implementation')
doc.add_paragraph('• FDSproject/backend/algorithms/FloydWarshallAlgorithm.java - Floyd-Warshall implementation')
doc.add_paragraph('• FDSproject/backend/algorithms/AlgorithmRunner.java - Java-Python integration')
doc.add_paragraph('• FDSproject/backend/app.py - Python Flask backend with hashtable implementation')

# Save the document
doc.save('c:/Users/ACER/Documents/impleFDS/DSA_Algorithms_and_Hashtable_Documentation.docx')
print('Word document created successfully!')
